from pathlib import Path
import re
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

workspace = Path.cwd()
parent = workspace.parent
output_dir = workspace / 'tools' / 'role_flow_complete_screens'
output_dir.mkdir(parents=True, exist_ok=True)

def make_image(title, lines, out_path: Path):
    w, h = 1200, 600
    img = Image.new('RGB', (w, h), 'white')
    draw = ImageDraw.Draw(img)
    try:
        font_b = ImageFont.truetype('arialbd.ttf', 28)
        font = ImageFont.truetype('arial.ttf', 18)
    except Exception:
        font_b = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.rectangle([(0,0),(w-1,70)], fill=(18,54,92))
    draw.text((24,18), title, fill='white', font=font_b)
    y = 90
    for l in lines:
        # wrap long lines
        if len(l) > 120:
            parts = [l[i:i+120] for i in range(0, len(l), 120)]
            for p in parts:
                draw.text((36,y), u"\u2022 "+p, fill='black', font=font)
                y += 28
        else:
            draw.text((36,y), u"\u2022 "+l, fill='black', font=font)
            y += 28
        if y > h-40:
            break
    img.save(out_path)

# Detailed per-role flows inferred from frontend code
flows = {
    'Common (App Start -> Logout)': [
        'SplashScreen (lib/Screens/StartUp-Screens/splash_screen.dart) — plays video, then -> OnboardingScreen',
        'OnboardingScreen -> LoginScreen (option to Skip -> Login)',
        'LoginScreen: submit credentials -> backend auth -> save tokens (AuthSession) -> role normalization',
        'Optional ChangePasswordScreen if OTC / first-login or password recovery flows',
        'PushNotificationService register and then navigate to role-specific Dashboard (RoleDashboard logic)',
        'Drawer / bottom navigation access modules, each module opens screens and detail flows',
        'Profile/Settings -> Change Password -> Logout (pushAndRemoveUntil LoginScreen)',
    ],

    'SuperAdmin': [
        'Login -> SuperAdminDashboard (SuperAdminDashboard widget)',
        'Access: Organization management, create admins (Create Admins screen), system audit logs, critical alerts popup',
        'User provisioning flows -> Create Admin -> confirmation -> notifications',
        'Reports & exports -> open report detail -> download PDF',
        'Logout via drawer -> confirmation dialog -> LoginScreen',
    ],

    'CEO / MD / ED (Leadership)': [
        'Login -> CeoDashboard / MdDashboard / ExecutiveDirectorDashboard',
        'Tabs: Home, People, Approvals, Client Visits, More',
        'Home: metrics, revenue, attendance; quick actions open Directory/Reports/Approvals',
        'People: open directory -> open employee profile -> view personal information -> open employee detail screens',
        'Approvals: list -> open approval detail -> approve/reject -> may open EmployeeApprovalDetailScreen or leave request flows',
        'Client Visits: open ClientVisitDashboard -> view visit -> open details -> verify / approve / view route and attachments',
        'More: Reports, Payroll Overview, Document Center -> open PDFs and exports',
        'Logout via drawer -> confirm -> LoginScreen',
    ],

    'Admin': [
        'Login -> AdminDashboard (Admin bottom nav with tabs)',
        'Tabs: Dashboard, Employees, Attendance, Leave, Meetings',
        'Employees: list -> add/edit employee -> registration flow (RegisterScreen) -> upload documents',
        'Attendance: view logs -> override check-ins/check-outs -> apply policies',
        'Leave: view requests -> approve/reject -> multi-stage approvals',
        'Meetings: schedule/manage meetings',
        'Admin can switch role to Employee (push EmployeeDashboard), or logout',
    ],

    'HR': [
        'Login -> HrDashboard',
        'HR tasks: onboarding approvals, employee documents review (EmployeeDocumentsScreen), leave and payroll approvals',
        'Employee onboarding: review RegisterScreen submissions -> accept/reject -> trigger notifications',
        'Document center: view and request re-uploads',
        'Logout via drawer -> confirm',
    ],

    'TL (Team Lead)': [
        'Login -> TLDashboard',
        'TL views team attendance, pending approvals and client journeys',
        'Client Visit approvals: when a visit is pending TL approval, open detail -> approve/return for edits',
        'Team journeys: view live journeys (ClientJourney screens) and verify locations',
        'Logout via drawer',
    ],

    'Employee': [
        'Login -> EmployeeDashboard',
        'Common features: Attendance (check-in/out), Tasks, Payslip, Documents, Approvals',
        'Start Client Visit: open ClientVisitDashboard -> Create Visit (ClientVisitCreateScreen) -> submit',
        'During a visit: status transitions -> travelling -> travel progress -> check-in (client selfie) -> in_progress -> complete -> summary',
        'Uploads: attachments, expenses via ClientVisitService.uploadFiles',
        'Start Client Journey: CreateClientJourneyScreen -> Start -> JourneyTracker starts foreground service -> location stream -> sync -> stop -> summary',
        'Profile: View/Edit personal info, upload documents via EmployeeDocumentsScreen',
        'Logout via drawer -> confirm',
    ],

    'ClientVisit Module (detailed flow)': [
        'Open ClientVisitDashboard (role-specific entry points)',
        'Create Visit: fill client, purpose, address, assign TL -> submit -> backend creates visit',
        'State machine for a visit (employee flow):',
        '  • Approved -> Employee starts travel (status=travelling)',
        '  • Travel progress -> if reachedClientAt == null show TravelProgressScreen',
        '  • On reach -> CheckInScreen (capture selfie, GPS) -> mark as in_progress',
        '  • In progress -> active visit screens (forms, attachments, expenses)',
        '  • Complete -> Summary screen with route, attachments, signatures',
        'Reviewer/Manager flow: list pending -> open detail -> approve/verify -> may mark TL/HR approvals',
        'Attachments: upload photos, receipts; server endpoint /attachments/',
    ],

    'ClientJourney (tracking) Module': [
        'Create Client Journey -> set client, purpose, destination, scheduled time',
        'Start Journey -> JourneyTracker.start() calls FlutterForegroundTask.startService',
        'Foreground service records GPS positions periodically, enqueues to SQLite queue, attempts sync via JourneyRepository.syncPending',
        'Live tracking: connect websocket via JourneyRepository.connect for live map',
        'Stop Journey -> stop service -> finalize route -> show summary and route report',
    ],
}

# Build document
print('Building document...')
doc = Document()
doc.add_heading('HRMS-BB — Complete Role & Module Flows', level=1)
doc.add_paragraph('This document describes end-to-end screen flows for each role from app start (splash) through module actions to logout. Generated from frontend code analysis.',)

for title, steps in flows.items():
    doc.add_heading(title, level=2)
    for i, s in enumerate(steps, start=1):
        doc.add_paragraph(f"{i}. {s}")
    # create and insert image
    # sanitize title for filename
    safe = re.sub(r'[^A-Za-z0-9._-]', '_', title)
    img_path = output_dir / f"{safe}.png"
    make_image(title, steps, img_path)
    doc.add_picture(str(img_path), width=Inches(6))

# Save
parent_doc = parent / 'HRMS-BB-Complete-Role-Flows.docx'
workspace_doc = workspace / 'HRMS-BB-Complete-Role-Flows.docx'
doc.save(str(parent_doc))
doc.save(str(workspace_doc))
print('Saved:', parent_doc)
print('Also saved:', workspace_doc)
print('Screenshots:', output_dir)
