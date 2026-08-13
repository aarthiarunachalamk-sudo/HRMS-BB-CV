from pathlib import Path
from docx import Document
from docx.shared import Inches

workspace = Path.cwd()
parent = workspace.parent
images_dir = workspace / 'tools' / 'indetailed_screens'

doc = Document()
doc.add_heading('HRMS-BB — App Flow (Brief)', level=1)
doc.add_paragraph('A concise summary of the main app flows and representative screenshots for quick review. Suitable for non-technical stakeholders.',)

sections = [
    ('Login & Onboarding',
     'Splash -> Onboarding -> Login. Users authenticate and are routed to role-specific dashboards.' ,
     'Common__App_Start_-__Logout__step01.png'),

    ('Employee — Attendance & Client Visit',
     'Check-in/out via Attendance tab (selfie + GPS). Create and run Client Visits with check-in at client location and submit attachments.',
     'Employee___Attendance___Client_Visit_Flow_step01.png'),

    ('Admin — Onboarding & Employee Management',
     'Add or review employees, upload documents, assign roles and payroll settings.',
     'Admin___Employee_Onboarding_Flow_step01.png'),

    ('Client Journey (Live Tracking)',
     'Create journey -> start foreground tracker -> live GPS capture and route summary on stop.',
     'ClientJourney__Tracking__Module_Detailed_step02.png'),
]

for title, desc, imgname in sections:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)
    img_path = images_dir / imgname
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5))
    else:
        doc.add_paragraph('(screenshot not available: {})'.format(imgname))

doc.add_heading('Quick Notes', level=2)
doc.add_paragraph('- If reviewers need more detail, see the full flow document attached separately.')
doc.add_paragraph('- For device installs, ensure older app versions are uninstalled or build versionCode is bumped.')

parent_doc = parent / 'HRMS-BB-App-Flow-Brief.docx'
workspace_doc = workspace / 'HRMS-BB-App-Flow-Brief.docx'
doc.save(str(parent_doc))
doc.save(str(workspace_doc))

print('Saved:', parent_doc)
print('Workspace copy:', workspace_doc)
