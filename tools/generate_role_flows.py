import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

workspace = Path.cwd()
parent = workspace.parent
output_dir = workspace / 'tools' / 'role_flow_screens'
output_dir.mkdir(parents=True, exist_ok=True)

def make_image(title, lines, out_path: Path):
    w, h = 1100, 420
    img = Image.new('RGB', (w, h), 'white')
    draw = ImageDraw.Draw(img)
    try:
        font_b = ImageFont.truetype('arialbd.ttf', 26)
        font = ImageFont.truetype('arial.ttf', 18)
    except Exception:
        font_b = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.rectangle([(0,0),(w-1,60)], fill=(30,60,90))
    draw.text((20,12), title, fill='white', font=font_b)
    y = 80
    for l in lines:
        draw.text((28,y), u"\u2022 "+l, fill='black', font=font)
        y += 28
    img.save(out_path)

roles = [
    'SuperAdmin','CEO','MD','ED','Admin','TL','HR','Employee','Finance','Manager','IT','Marketing','ClientVisit','ClientJourney'
]

def common_flow():
    return [
        'App launch: Splash screen (video) -> Onboarding',
        'Login screen: enter Employee code + password; Forgot password path',
        'OTC/Change password flow if required',
        'Auth tokens saved (AuthSession)',
        'Push notifications registration',
        'Navigate to role-specific dashboard',
        'Use drawer or bottom navigation to access modules',
        'Open module -> perform actions -> save/submit',
        'Open Profile -> Change Password -> Logout',
    ]

role_specific = {
    'SuperAdmin': [
        'Access SuperAdmin dashboard',
        'Manage organizations, create admins',
        'User provisioning & critical alerts',
        'View system-wide reports and audit logs',
    ],
    'CEO': [
        'Open CEO dashboard: Home / People / Approvals / Client Visits / More',
        'View revenue, attendance, pending approvals',
        'Open approval details -> Approve/Reject',
        'Open reports -> download or view PDF',
    ],
    'MD': ['MD dashboard: metrics, reports, approvals (similar to CEO)'],
    'ED': ['Executive Director dashboard and quick approvals'],
    'Admin': [
        'Admin bottom nav: Dashboard, Employees, Attendance, Leave, Meetings',
        'Create/manage employees, run attendance/checkout actions',
        'Open Admin flow screens for bulk ops',
    ],
    'TL': [
        'Team Lead dashboard: team attendance, approvals, journeys',
        'Approve client visits, verify check-ins',
    ],
    'HR': [
        'HR dashboard: employee profiles, leave approvals, onboarding docs',
        'Final approvals for client visits when required',
    ],
    'Employee': [
        'Employee dashboard: personal overview, attendance, tasks',
        'Start client visits / journeys, check-in, upload selfies and receipts',
        'View payslip, documents, submit requests',
    ],
    'Finance': ['Payroll overview, payslip access, finance reports'],
    'Manager': ['Manager dashboard: tasks, team overview, approvals'],
    'IT': ['IT dashboard: tasks, system notifications, support tickets'],
    'Marketing': ['Marketing dashboard: tasks and campaigns'],
    'ClientVisit': [
        'Create Client Visit -> Fill details -> Submit',
        'Employee: Travel -> Check-in -> Active -> Complete -> Summary',
        'Manager/TL/SuperAdmin: Review -> Approve/Reject -> Verify (signature, GPS, attachments)',
        'Uploads: attachments, expenses, signatures',
    ],
    'ClientJourney': [
        'Create Journey -> Start tracking (foreground service)',
        'ForegroundTask saves location -> sync to server -> live map',
        'Stop journey -> complete summary and route report',
    ],
}

# Build the document
doc = Document()
doc.add_heading('HRMS-BB — Detailed Role Flows', level=1)

doc.add_heading('Common App Flow (applies to all roles)', level=2)
for p in common_flow():
    doc.add_paragraph(p)

for role in roles:
    doc.add_heading(role, level=2)
    steps = common_flow()
    specific = role_specific.get(role, [])
    for s in specific:
        steps.append(s)
    # write steps
    for s in steps:
        doc.add_paragraph(s)
    # make and add image
    img_path = output_dir / f"{role}_flow.png"
    make_image(f"{role} Flow", steps[:8] + (specific if specific else []), img_path)
    doc.add_picture(str(img_path), width=Inches(6))

# Save both to parent and workspace
parent_doc = parent / 'HRMS-BB-Role-Flows.docx'
workspace_doc = workspace / 'HRMS-BB-Role-Flows.docx'
doc.save(str(parent_doc))
doc.save(str(workspace_doc))
print('Saved:', parent_doc)
print('Also saved:', workspace_doc)
print('Screenshots:', output_dir)
