import os
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

# Paths
workspace_dir = Path.cwd()
parent_dir = workspace_dir.parent
output_dir = workspace_dir / "tools" / "module_screenshots"
output_dir.mkdir(parents=True, exist_ok=True)

# Detect modules
modules = []
# top-level frontend and backend
if (workspace_dir / 'backend').exists():
    modules.append(('backend (project)', workspace_dir / 'backend'))
if (workspace_dir / 'frontend').exists():
    modules.append(('frontend (app)', workspace_dir / 'frontend'))

# backend apps
backend_apps_dir = workspace_dir / 'backend'
if backend_apps_dir.exists():
    for entry in sorted(backend_apps_dir.iterdir()):
        if entry.is_dir() and entry.name not in ('.venv', '__pycache__', 'media'):
            # skip project subfolder named 'backend'
            if entry.name == 'backend':
                modules.append(('backend (django settings)', entry))
                continue
            modules.append((entry.name, entry))

# Heuristic flow builder
def build_flow(name, path: Path):
    steps = []
    if (path / 'models.py').exists() or any(p.suffix=='.py' for p in path.glob('*')):
        steps.append('Data models: defined in models.py')
    if (path / 'serializers.py').exists():
        steps.append('Serializers: translate models to API payloads')
    if (path / 'views.py').exists() or (path / 'views').exists():
        steps.append('Views / Controllers: handle HTTP requests')
    if (path / 'urls.py').exists():
        steps.append('Routing: urls.py maps endpoints to views')
    if (path / 'consumers.py').exists() or (path / 'routing.py').exists():
        steps.append('WebSockets/Realtime: consumers + routing')
    if (path / 'storage.py').exists():
        steps.append('Storage: file/photo storage handling')
    if (path / 'management').exists():
        steps.append('Management commands: custom CLI tasks')
    if (path / 'tests.py').exists() or list(path.glob('test_*.py')):
        steps.append('Tests: unit/integration tests present')
    if not steps:
        steps = ['Inspect files for specific flows (models/views/urls)']
    return steps

# Create visual "screenshot" images
def create_screenshot(name, steps, out_path: Path):
    width, height = 1000, 300
    img = Image.new('RGB', (width, height), color='white')
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype('arial.ttf', 20)
    except Exception:
        font = ImageFont.load_default()
    title = f"Module: {name}"
    draw.text((20, 20), title, fill='black', font=font)
    y = 60
    for s in steps:
        draw.text((40, y), f"- {s}", fill='black', font=font)
        y += 30
    img.save(out_path)

# Build document
doc = Document()
doc.add_heading('Module Working Flows', level=1)
for name, path in modules:
    doc.add_heading(name, level=2)
    steps = build_flow(name, path)
    for s in steps:
        doc.add_paragraph(s)
    img_name = f"{name.replace(' ', '_')}.png"
    img_path = output_dir / img_name
    create_screenshot(name, steps, img_path)
    # Add image to doc
    doc.add_picture(str(img_path), width=Inches(6))

# Save to parent directory and workspace root for convenience
parent_doc = parent_dir / 'HRMS-BB-Module-Flows.docx'
workspace_doc = workspace_dir / 'HRMS-BB-Module-Flows.docx'
doc.save(str(parent_doc))
doc.save(str(workspace_doc))

print('Saved:', parent_doc)
print('Also saved:', workspace_doc)
print('Screenshots folder:', output_dir)
