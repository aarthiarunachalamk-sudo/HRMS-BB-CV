from pathlib import Path
import re
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

workspace = Path.cwd()
parent = workspace.parent
output_dir = workspace / 'tools' / 'indetailed_screens'
output_dir.mkdir(parents=True, exist_ok=True)

def make_image(title, subtitle, text, out_path: Path):
    w, h = 1200, 700
    img = Image.new('RGB', (w, h), 'white')
    draw = ImageDraw.Draw(img)
    try:
        font_b = ImageFont.truetype('arialbd.ttf', 26)
        font = ImageFont.truetype('arial.ttf', 16)
    except Exception:
        font_b = ImageFont.load_default()
        font = ImageFont.load_default()
    draw.rectangle([(0,0),(w-1,70)], fill=(18,54,92))
    draw.text((24,18), title, fill='white', font=font_b)
    draw.text((24,78), subtitle, fill='black', font=font)
    y = 110
    lines = []
    for l in text.split('\n'):
        if len(l) > 120:
            parts = [l[i:i+120] for i in range(0, len(l), 120)]
            lines.extend(parts)
        else:
            lines.append(l)
    for l in lines:
        draw.text((36,y), u"\u2022 "+l, fill='black', font=font)
        y += 24
        if y > h-40:
            break
    img.save(out_path)


flows = {
    'Common (App Start -> Logout)': [
        'SplashScreen — plays intro video and performs cold-start checks',
        'OnboardingScreen — slides explaining app features (may be skipped)',
        'LoginScreen — enter credentials, two-factor if enabled; on success persist tokens',
        'Home/Dashboard — role-specific widgets and quick actions',
        'Common Drawer/Profile — settings, change password, logout',
    ],

    'Employee — Attendance & Client Visit Flow': [
        'Open Attendance tab -> Check-in button -> capture geolocation and selfie -> send to /attendance/checkin/',
        'Check-in success -> show todays summary -> check-out button appears after time threshold',
        'Open Client Visits -> Create Visit -> fill client, purpose, address -> submit -> status=PENDING',
        'Start Visit -> navigation starts -> Journey Tracker runs (foreground service) -> location sync every N seconds',
        'On reaching client -> CheckInScreen -> capture selfie & GPS -> status=IN_PROGRESS -> fill visit forms',
        'Complete Visit -> Submit summary -> generate report and attachments -> status=COMPLETED',
    ],

    'Admin — Employee Onboarding Flow': [
        'Open Employees -> Add Employee -> fill personal & company details -> upload documents',
        'Submit -> backend creates user -> sends invite email',
        'New user appears in Employees list -> click to review documents -> approve/reject',
        'Assign roles/permissions -> set payroll and reporting manager -> notify employee',
    ],

    'ClientJourney (Tracking) Module Detailed': [
        'Create Journey -> set destination, client, scheduled time -> save -> status=SCHEDULED',
        'Start Journey -> foreground service start -> show map with live polyline',
        'During journey -> periodic GPS capture -> local queue -> attempt sync to /journeys/sync/',
        'Pause/Resume -> user can pause tracking (stop sending but store locally) -> resume continues',
        'Stop Journey -> finalize route -> show summary with distance/time -> exportable report',
    ],
}

print('Building in-detailed flows document...')
doc = Document()
doc.add_heading('HRMS-BB — In-Depth Screen Flows', level=1)
doc.add_paragraph('This document contains per-screen detailed steps and illustrative screenshots generated from repository analysis.',)

for title, steps in flows.items():
    doc.add_heading(title, level=2)
    for idx, step in enumerate(steps, start=1):
        # split a short subtitle and details if long
        parts = step.split('->')
        subtitle = parts[0].strip()
        detail = step
        doc.add_heading(f"{idx}. {subtitle}", level=3)
        doc.add_paragraph(detail)
        safe_title = re.sub(r'[^A-Za-z0-9._-]', '_', title)
        img_name = f"{safe_title}_step{idx:02d}.png"
        img_path = output_dir / img_name
        make_image(title, f"Step {idx}: {subtitle}", detail, img_path)
        doc.add_picture(str(img_path), width=Inches(6))

parent_doc = parent / 'HRMS-BB-InDetailed-Flows.docx'
workspace_doc = workspace / 'HRMS-BB-InDetailed-Flows.docx'
doc.save(str(parent_doc))
doc.save(str(workspace_doc))
print('Saved:', parent_doc)
print('Also saved:', workspace_doc)
print('Screenshots:', output_dir)
